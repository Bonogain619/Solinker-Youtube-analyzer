import streamlit as st
from googleapiclient.discovery import build
import pandas as pd
import plotly.express as px
import isodate
import requests
import json
from datetime import datetime
import io 
from docx import Document 
from docx.oxml.ns import qn 
from docx.shared import Pt # í°íŠ¸ í¬ê¸° ì¡°ì ˆìš©

# ==============================================================================
# [í•„ìŠ¹ ì„¤ì •] check.pyì—ì„œ ì„±ê³µí–ˆë˜ ê·¸ 'Gemini í‚¤'ë¥¼ ë”°ì˜´í‘œ ì•ˆì— ë¶™ì—¬ë„£ìœ¼ì„¸ìš”!
# ==============================================================================
GEMINI_API_KEY = "AIzaSyASndV5RssUI9Uj3Shuwm8BdhnFP5F7vx4"
# ==============================================================================

st.set_page_config(page_title="Solinker Channel Analyzer", page_icon="âš¡", layout="wide")

# [ìŠ¤íƒ€ì¼] UI ìµœì í™”
st.markdown("""
<style>
    .report-box {
        border: 2px solid #e0e0e0;
        padding: 30px;
        border-radius: 15px;
        background-color: #f9f9f9;
        color: #333333;
        font-size: 1.2rem !important; 
        line-height: 1.8 !important;
        margin-bottom: 30px;
    }
    .stTabs [data-baseweb="tab"] {
        font-size: 1.1rem;
        font-weight: 600;
    }
    .big-text {
        font-size: 1.5rem !important;
        font-weight: bold;
        margin-bottom: 10px;
        color: #1E1E1E;
    }
</style>
""", unsafe_allow_html=True)

# 1. ì´ˆê¸°í™”
if "run_pro" not in st.session_state: st.session_state.run_pro = False
if "messages" not in st.session_state: st.session_state.messages = []
if "data" not in st.session_state: st.session_state.data = None

# 2. ì‚¬ì´ë“œë°” UI
with st.sidebar:
    st.header("ğŸ”§ ì„¤ì • íŒ¨ë„")
    
    if GEMINI_API_KEY.startswith("AIza"):
        st.success(f"âœ… AI ì—”ì§„ ì¤€ë¹„ ì™„ë£Œ")
    else:
        st.error("ğŸš¨ ì½”ë“œ 12ë²ˆì§¸ ì¤„ì— Gemini í‚¤ë¥¼ ë„£ì–´ì£¼ì„¸ìš”!")

    with st.expander("ğŸ”‘ ìœ íŠœë¸Œ í‚¤ ì…ë ¥", expanded=True):
        yt_key = st.text_input("YouTube API Key", type="password")
    
    st.divider()
    handle_input = st.text_input("ì±„ë„ í•¸ë“¤ (@í¬í•¨)", placeholder="@í•¸ë“¤ëª…")
    
    if st.button("âš¡ ì‹¬ì¸µ ë¶„ì„ ì‹œì‘", type="primary"):
        st.session_state.run_pro = True
        st.session_state.messages = [] 
        st.session_state.data = None

# 3. ìœ í‹¸ë¦¬í‹° í•¨ìˆ˜ë“¤
def get_youtube(api_key): return build("youtube", "v3", developerKey=api_key)

def check_is_shorts(video_id):
    try: return requests.head(f"https://www.youtube.com/shorts/{video_id}", allow_redirects=False, timeout=2).status_code == 200
    except: return False

def format_duration(seconds):
    m, s = divmod(int(seconds), 60)
    h, m = divmod(m, 60)
    if h > 0: return f"{h}:{m:02d}:{s:02d}"
    return f"{m}:{s:02d}"

def get_channel_stats(yt, handle):
    try:
        res = yt.search().list(part="id,snippet", q=handle, type="channel", maxResults=1).execute()
        if not res["items"]: return None
        ch_id = res["items"][0]["id"]["channelId"]
        item = yt.channels().list(part="statistics,contentDetails,snippet", id=ch_id).execute()["items"][0]
        return {
            "title": item["snippet"]["title"],
            "thumbnail": item["snippet"]["thumbnails"]["high"]["url"],
            "subs": int(item["statistics"]["subscriberCount"]),
            "views": int(item["statistics"]["viewCount"]),
            "video_count": int(item["statistics"]["videoCount"]),
            "upload_id": item["contentDetails"]["relatedPlaylists"]["uploads"],
            "desc": item["snippet"]["description"]
        }
    except Exception as e: st.error(f"ì±„ë„ ê²€ìƒ‰ ì‹¤íŒ¨: {e}"); return None

def get_recent_videos(yt, upload_id):
    try:
        res = yt.playlistItems().list(part="snippet,contentDetails", playlistId=upload_id, maxResults=50).execute()
        vid_ids = [i["contentDetails"]["videoId"] for i in res["items"]]
        vid_res = yt.videos().list(part="statistics,contentDetails,snippet", id=",".join(vid_ids)).execute()
        videos = []
        status_text = st.empty()
        total = len(vid_res["items"])
        for i, item in enumerate(vid_res["items"]):
            stats = item["statistics"]
            dur = isodate.parse_duration(item["contentDetails"]["duration"]).total_seconds()
            is_s = False
            if dur <= 180:
                status_text.caption(f"ğŸ” ì˜ìƒ ë¶„ì„ ì¤‘ ({i+1}/{total})...")
                if check_is_shorts(item['id']): is_s = True
            videos.append({
                "ì œëª©": item["snippet"]["title"],
                "ì¡°íšŒìˆ˜": int(stats.get("viewCount", 0)),
                "ì¢‹ì•„ìš”": int(stats.get("likeCount", 0)),
                "ëŒ“ê¸€": int(stats.get("commentCount", 0)),
                "ê¸¸ì´": format_duration(dur),
                "ë‚ ì§œ": item["snippet"]["publishedAt"][:10],
                "ìœ í˜•": "Shorts" if is_s else "Long-form"
            })
        status_text.empty()
        return pd.DataFrame(videos)
    except: return pd.DataFrame()

# [í•µì‹¬] ì›Œë“œ ìƒì„±ê¸° (í‘œ ì¸ì‹ ê¸°ëŠ¥ ì¶”ê°€ë¨)
def create_docx(text, title="ë¬¸ì„œ"):
    doc = Document()
    
    # 1. ë¬¸ì„œ ê¸°ë³¸ í°íŠ¸ ì„¤ì • (í•œê¸€ ê¹¨ì§ ë°©ì§€)
    style = doc.styles['Normal']
    style.font.name = 'Malgun Gothic'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    
    doc.add_heading(title, 0)
    
    # ì¤„ ë‹¨ìœ„ë¡œ ì²˜ë¦¬í•˜ë˜, í‘œ(Table)ë¥¼ ë§Œë‚˜ë©´ ëª¨ì•„ì„œ í•œ ë²ˆì— ì²˜ë¦¬
    lines = text.split('\n')
    table_buffer = [] # í‘œ ë‚´ìš©ì„ ì„ì‹œ ì €ì¥í•  ê³µê°„
    
    for line in lines:
        line = line.strip()
        
        # (1) í‘œê°€ ì‹œì‘ë˜ê±°ë‚˜ ì´ì–´ì§€ëŠ” ê²½ìš° (|ë¡œ ì‹œì‘)
        if line.startswith('|'):
            table_buffer.append(line)
        else:
            # (2) í‘œê°€ ëë‚¬ëŠ”ë° ë²„í¼ì— ë‚´ìš©ì´ ìˆë‹¤ë©´ -> í‘œ ìƒì„±!
            if table_buffer:
                _add_table_to_doc(doc, table_buffer)
                table_buffer = [] # ë²„í¼ ì´ˆê¸°í™”
            
            # (3) ì¼ë°˜ í…ìŠ¤íŠ¸ ì²˜ë¦¬
            if not line: continue
            
            if line.startswith('### '): doc.add_heading(line.replace('### ', ''), level=3)
            elif line.startswith('## '): doc.add_heading(line.replace('## ', ''), level=2)
            elif line.startswith('# '): doc.add_heading(line.replace('# ', ''), level=1)
            elif line.startswith('- ') or line.startswith('* '): doc.add_paragraph(line, style='List Bullet')
            else: doc.add_paragraph(line)
            
    # ë°˜ë³µë¬¸ì´ ëë‚¬ëŠ”ë° ë§ˆì§€ë§‰ì— í‘œê°€ ë‚¨ì•„ìˆì„ ê²½ìš° ì²˜ë¦¬
    if table_buffer:
        _add_table_to_doc(doc, table_buffer)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ë‚´ë¶€ í•¨ìˆ˜: ë§ˆí¬ë‹¤ìš´ í‘œë¥¼ ì›Œë“œ í‘œë¡œ ë³€í™˜
def _add_table_to_doc(doc, markdown_lines):
    # ë°ì´í„° íŒŒì‹±
    rows = []
    for line in markdown_lines:
        # | êµ¬ë¶„ìë¡œ ìë¥´ê³  ì•ë’¤ ê³µë°± ì œê±°
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)
    
    # êµ¬ë¶„ì„ (---|---) ì œê±°: ë³´í†µ ë‘ ë²ˆì§¸ ì¤„ì— ìˆìŒ
    real_rows = [r for r in rows if not set(''.join(r)).issubset(set('-:| '))]
    
    if not real_rows: return

    # í‘œ ìƒì„±
    num_cols = len(real_rows[0])
    table = doc.add_table(rows=len(real_rows), cols=num_cols)
    table.style = 'Table Grid' # ê²©ì ë¬´ëŠ¬ ìŠ¤íƒ€ì¼ ì ìš©
    
    # ë°ì´í„° ì±„ìš°ê¸° & í°íŠ¸ ì„¤ì •
    for i, row_data in enumerate(real_rows):
        row = table.rows[i]
        for j, text in enumerate(row_data):
            if j < len(row.cells):
                cell = row.cells[j]
                cell.text = text
                # í‘œ ì•ˆì˜ ê¸€ì”¨ë„ 'ë§‘ì€ ê³ ë”•' ê°•ì œ ì ìš©
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Malgun Gothic'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

# 4. AI ì—°ê²°
def call_gemini_rest(prompt):
    models = ["gemini-flash-latest", "gemini-1.5-flash", "gemini-pro"]
    for model in models:
        try:
            url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={GEMINI_API_KEY}"
            resp = requests.post(url, headers={'Content-Type': 'application/json'}, json={"contents": [{"parts": [{"text": prompt}]}]}, timeout=30)
            if resp.status_code == 200: return resp.json()['candidates'][0]['content']['parts'][0]['text']
        except: continue
    return "âŒ AI ì—°ê²° ì‹¤íŒ¨"

def generate_pro_insight(channel, df):
    prompt = f"""
    ë‹¹ì‹ ì€ ìµœì‹  íŠ¸ë Œë“œë¥¼ ë°˜ì˜í•˜ëŠ” ìœ íŠœë¸Œ ì»¨ì„¤í„´íŠ¸ì…ë‹ˆë‹¤. (ê¸°ì¤€ì¼: {datetime.now().strftime('%Y-%m-%d')})
    
    [ì •ì±… ê°€ì´ë“œ]
    - ì‡¼ì¸ ëŠ” ìµœëŒ€ 3ë¶„ê¹Œì§€ ê°€ëŠ¥í•©ë‹ˆë‹¤. 60ì´ˆ ì œí•œì´ë¼ê³  ë‹¨ì • ì§“ì§€ ë§ˆì„¸ìš”.
    
    [ì±„ë„ ì •ë³´]
    - ì±„ë„ëª…: {channel['title']}
    - êµ¬ë…ì: {channel['subs']}ëª…
    
    [ë°ì´í„°]
    {df[['ìœ í˜•', 'ì œëª©', 'ì¡°íšŒìˆ˜', 'ì¢‹ì•„ìš”', 'ëŒ“ê¸€', 'ê¸¸ì´', 'ë‚ ì§œ']].to_string(index=False)}
    
    ìœ„ ë°ì´í„°ë¥¼ ë°”íƒ•ìœ¼ë¡œ ë‹¤ìŒ ë‚´ìš©ì„ ë§ˆí¬ë‹¤ìš´ìœ¼ë¡œ ì‘ì„±í•˜ì„¸ìš”:
    1. ğŸ“Š íŒ©íŠ¸ ì²´í¬ (ì¡°íšŒìˆ˜ ë° ì¶©ì„±ë„ ë¶„ì„) - **ë°˜ë“œì‹œ í‘œ(Table)ë¥¼ ì‚¬ìš©í•˜ì—¬ ë°ì´í„°ë¥¼ ë¹„êµí•˜ì„¸ìš”.**
    2. ğŸš¨ ëƒ‰ì •í•œ ë¹„íŒ (ì„±ì¥ ì •ì²´ ì›ì¸)
    3. ğŸš€ ì†”ë£¨ì…˜ 3ê°€ì§€ (êµ¬ì²´ì  ì‹¤í–‰ ë°©ì•ˆ)
    """
    return call_gemini_rest(prompt)

def ask_gemini_chat(question, context_report):
    prompt = f"ë‹¹ì‹ ì€ ìœ íŠœë¸Œ ì»¨ì„¤í„´íŠ¸ì…ë‹ˆë‹¤.\n[ë¦¬í¬íŠ¸]\n{context_report}\n[ì§ˆë¬¸]\n{question}\në‹µë³€í•´ì£¼ì„¸ìš”."
    return call_gemini_rest(prompt)

# 5. ë©”ì¸ ì‹¤í–‰
if st.session_state.run_pro and yt_key:
    yt = get_youtube(yt_key)
    stats = get_channel_stats(yt, handle_input)
    if stats:
        with st.spinner("ë°ì´í„° ë¶„ì„ ì¤‘..."):
            df = get_recent_videos(yt, stats["upload_id"])
            if not df.empty:
                report = generate_pro_insight(stats, df)
                st.session_state.data = (stats, df, report)
                st.session_state.run_pro = False 
                st.rerun()

# 6. í™”ë©´ ì¶œë ¥
if st.session_state.data is not None:
    stats, df, report = st.session_state.data
    
    c1, c2 = st.columns([1, 6])
    with c1: st.image(stats["thumbnail"], width=100)
    with c2: st.title(stats["title"])
    st.divider()
    
    t1, t2 = st.tabs(["ğŸ“„ AI ì‹¬ì¸µ ë¦¬í¬íŠ¸ & ì±„íŒ…", "ğŸ“ˆ ë°ì´í„° ìƒì„¸"])
    
    with t1: 
        st.markdown(f'<div class="report-box">{report}</div>', unsafe_allow_html=True)
        st.divider()
        st.subheader("ğŸ’¬ AI ì»¨ì„¤í„´íŠ¸ì—ê²Œ ì§ˆë¬¸í•˜ê¸°")
        
        for msg in st.session_state.messages:
            with st.chat_message("user" if msg['role']=="user" else "assistant"):
                st.markdown(msg['content'])

        if prompt := st.chat_input("ì§ˆë¬¸ì„ ì…ë ¥í•˜ì„¸ìš”..."):
            st.session_state.messages.append({"role": "user", "content": prompt})
            with st.chat_message("user"): st.markdown(prompt)
            with st.chat_message("assistant"):
                with st.spinner("ìƒê° ì¤‘..."):
                    ans = ask_gemini_chat(prompt, report)
                    st.markdown(ans)
                    st.session_state.messages.append({"role": "assistant", "content": ans})
        
    with t2:
        st.markdown('<p class="big-text">ğŸ“Š ìƒì„¸ ë°ì´í„° í…Œì´ë¸”</p>', unsafe_allow_html=True)
        st.dataframe(
            df[['ë‚ ì§œ', 'ìœ í˜•', 'ì œëª©', 'ê¸¸ì´', 'ì¡°íšŒìˆ˜', 'ì¢‹ì•„ìš”', 'ëŒ“ê¸€']],
            use_container_width=True,
            hide_index=True,
            column_config={
                "ì¡°íšŒìˆ˜": st.column_config.NumberColumn(format="%d"),
                "ì¢‹ì•„ìš”": st.column_config.NumberColumn(format="%d"),
                "ëŒ“ê¸€": st.column_config.NumberColumn(format="%d")
            }
        )
        st.markdown('<p class="big-text">ğŸ”´ Shorts vs ğŸ”µ Long-form ì„±ê³¼ ë¹„êµ</p>', unsafe_allow_html=True)
        fig = px.scatter(
            df, x="ë‚ ì§œ", y="ì¡°íšŒìˆ˜", size="ì¢‹ì•„ìš”", color="ìœ í˜•",
            color_discrete_map={"Shorts": "#FF4B4B", "Long-form": "#1C83E1"},
            hover_data=["ì œëª©", "ê¸¸ì´"], title="ì˜ìƒ ì„±ê³¼ & ì¶©ì„±ë„ ë¶„í¬"
        )
        st.plotly_chart(fig, use_container_width=True)

    with st.sidebar:
        st.divider()
        st.header("ğŸ“‚ ê²°ê³¼ ì €ì¥ì†Œ")
        file_prefix = f"Solinker_{stats['title']}_{datetime.now().strftime('%Y%m%d')}"
        
        docx_buffer = create_docx(report, title=f"{stats['title']} ì±„ë„ ë¶„ì„ ë¦¬í¬íŠ¸")
        st.download_button("ğŸ“„ ë¦¬í¬íŠ¸ ë‹¤ìš´ë¡œë“œ (.docx)", docx_buffer, f"{file_prefix}_Report.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        
        buffer = io.BytesIO()
        with pd.ExcelWriter(buffer, engine='xlsxwriter') as writer:
            df.to_excel(writer, index=False, sheet_name='Sheet1')
        st.download_button("ğŸ“Š ë°ì´í„° ë‹¤ìš´ë¡œë“œ (.xlsx)", buffer, f"{file_prefix}_Data.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        
        if st.session_state.messages:
            chat_full_text = ""
            for msg in st.session_state.messages:
                role = "ğŸ‘¤ ì§ˆë¬¸" if msg['role'] == "user" else "ğŸ¤– ë‹µë³€"
                chat_full_text += f"## {role}\n{msg['content']}\n\n"
            chat_docx = create_docx(chat_full_text, title=f"{stats['title']} AI ìƒë‹´ ê¸°ë¡")
            st.download_button("ğŸ’¬ ìƒë‹´ ê¸°ë¡ ë‹¤ìš´ë¡œë“œ (.docx)", chat_docx, f"{file_prefix}_Chat.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

else:
    st.title("ğŸ¥ Solinker Channel Analyzer")
    st.markdown("ì™¼ìª½ ì‚¬ì´ë“œë°”ì— **ìœ íŠœë¸Œ í‚¤**ì™€ **í•¸ë“¤**ì„ ì…ë ¥í•˜ê³  **[ì‹¬ì¸µ ë¶„ì„ ì‹œì‘]**ì„ ëˆŒëŸ¬ì£¼ì„¸ìš”.")